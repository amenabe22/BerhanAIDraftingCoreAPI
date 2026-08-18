#!/usr/bin/env python3
"""Run a bar/law exam (.docx or .doc) through /legal-search/stream and write answers only.

Usage:
  # API must be running (e.g. uvicorn on :8000)
  python examtest.py path/to/exam.docx
  python examtest.py path/to/exam.doc -o answers.md --format md
  python examtest.py exam.docx --language am --api http://127.0.0.1:8000

Uses POST /legal-search/stream (KB-grounded Qdrant answers). Override with --endpoint.
Requires: httpx, python-docx (already in requirements.txt)
Legacy .doc files are converted via macOS textutil or LibreOffice soffice.
"""

from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import sys
import tempfile
import uuid
from pathlib import Path

import httpx
from docx import Document


# Full exams can take several minutes of LLM + retrieval time.
DEFAULT_TIMEOUT_SECONDS = 900.0
SUPPORTED_SUFFIXES = {".docx", ".doc"}

PROMPT = """You are taking a law / bar examination.

Below is the full exam text extracted from a Word document.

Instructions:
1. Answer EVERY question in order (Q1, Q2, …).
2. For multiple choice: give the letter (A/B/C/D) first, then a short justification with article/proclamation cites when known.
3. For essays / short answers: give a concise model answer with pinpoint legal citations.
4. Output ANSWERS ONLY — no restating the full questions, no preamble, no closing remarks.
5. Use this format exactly:

## Q1
<answer>

## Q2
<answer>

---
EXAM TEXT:
"""

# Kept for --full-dump only. Prefer per-question QUESTION_PROMPT in exam_parse.


def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise SystemExit(f"No text found in {path}")
    return text


def _convert_doc_via_textutil(path: Path, dest_txt: Path) -> None:
    subprocess.run(
        ["textutil", "-convert", "txt", "-output", str(dest_txt), str(path)],
        check=True,
        capture_output=True,
        text=True,
    )


def _convert_doc_via_soffice(path: Path, out_dir: Path) -> Path:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) not found")
    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "txt:Text",
            "--outdir",
            str(out_dir),
            str(path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    converted = out_dir / f"{path.stem}.txt"
    if not converted.exists():
        # LibreOffice sometimes mangles names; pick the newest .txt
        txts = sorted(out_dir.glob("*.txt"), key=lambda p: p.stat().st_mtime, reverse=True)
        if not txts:
            raise RuntimeError(f"soffice produced no .txt for {path.name}")
        converted = txts[0]
    return converted


def extract_doc_text(path: Path) -> str:
    """Extract plain text from legacy .doc (OLE) Word files."""
    with tempfile.TemporaryDirectory(prefix="examtest_doc_") as tmp:
        tmp_dir = Path(tmp)
        dest_txt = tmp_dir / f"{path.stem}.txt"
        errors: list[str] = []

        if shutil.which("textutil"):
            try:
                _convert_doc_via_textutil(path, dest_txt)
                text = dest_txt.read_text(encoding="utf-8", errors="replace").strip()
                if text:
                    return text
                errors.append("textutil produced empty text")
            except (OSError, subprocess.CalledProcessError) as exc:
                errors.append(f"textutil failed: {exc}")

        try:
            converted = _convert_doc_via_soffice(path, tmp_dir)
            text = converted.read_text(encoding="utf-8", errors="replace").strip()
            if text:
                return text
            errors.append("soffice produced empty text")
        except (OSError, RuntimeError, subprocess.CalledProcessError) as exc:
            errors.append(f"soffice failed: {exc}")

        detail = "; ".join(errors) if errors else "no converter available"
        raise SystemExit(f"Could not extract text from {path}: {detail}")


def extract_exam_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        # Some misnamed .doc files are actually OOXML — try python-docx first.
        try:
            return extract_docx_text(path)
        except Exception:
            # Fall through only if it looks like legacy binary disguised as docx
            return extract_doc_text(path)
    if suffix == ".doc":
        # Sometimes .doc is actually OOXML
        try:
            return extract_docx_text(path)
        except Exception:
            return extract_doc_text(path)
    raise SystemExit(f"Unsupported file type: {suffix} (use .doc or .docx)")


def build_question_messages(exam_text: str) -> list[dict]:
    """Strip printed keys, split into questions, wrap each in QUESTION_PROMPT."""
    from app.services.legal.exam_parse import QUESTION_PROMPT, split_exam_questions, strip_printed_keys

    cleaned = strip_printed_keys(exam_text)
    questions = split_exam_questions(cleaned)
    return [
        {"n": q["n"], "stem": q["stem"], "message": QUESTION_PROMPT + q["text"]}
        for q in questions
    ]


def stream_exam(
    *,
    api: str,
    endpoint: str,
    message: str,
    language: str | None,
    model: str | None,
    enable_reasoning: bool,
    timeout: float,
) -> dict:
    path = endpoint if endpoint.startswith("/") else f"/{endpoint}"
    url = api.rstrip("/") + path
    thread_id = str(uuid.uuid4())
    payload: dict = {
        "message": message,
        "thread_id": thread_id,
        "enable_reasoning": enable_reasoning,
    }
    if language:
        payload["language"] = language
    if model:
        payload["model"] = model

    tokens: list[str] = []
    citations: list = []
    grounding: dict | None = None
    tool_queries: list[str] = []
    timeout_cfg = httpx.Timeout(timeout, connect=30.0)
    with httpx.Client(timeout=timeout_cfg) as client:
        with client.stream("POST", url, json=payload) as resp:
            if resp.status_code != 200:
                body = resp.read().decode("utf-8", errors="replace")
                raise SystemExit(f"HTTP {resp.status_code}: {body}")
            for line in resp.iter_lines():
                if not line.startswith("data: "):
                    continue
                raw = line[len("data: ") :]
                try:
                    evt = json.loads(raw)
                except json.JSONDecodeError:
                    continue
                et = evt.get("type")
                if et == "token":
                    chunk = evt.get("content") or ""
                    tokens.append(chunk)
                    print(chunk, end="", flush=True, file=sys.stderr)
                elif et == "status":
                    msg = evt.get("message") or ""
                    if msg:
                        print(f"\n[status] {msg}", file=sys.stderr)
                elif et == "citations":
                    citations = evt.get("citations") or evt.get("items") or citations
                elif et == "grounding":
                    grounding = evt
                elif et == "tool_query":
                    q = (evt.get("query") or "").strip()
                    if q:
                        tool_queries.append(q)
                elif et == "thread_id":
                    thread_id = evt.get("thread_id") or thread_id
                elif et == "error":
                    raise SystemExit(f"Agent error: {evt.get('message')}")
    print(file=sys.stderr)
    return {
        "thread_id": thread_id,
        "answer": "".join(tokens).strip(),
        "citations": citations,
        "grounding": grounding,
        "tool_queries": tool_queries,
    }


DEFAULT_ENDPOINT = "/legal-search/stream"


def main() -> None:
    p = argparse.ArgumentParser(description="Answer a .doc/.docx law exam via legal-search/stream.")
    p.add_argument("exam", type=Path, help="Path to exam .doc or .docx")
    p.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="Output path (default: <exam>_answers.md or .txt)",
    )
    p.add_argument(
        "--format",
        choices=("md", "txt"),
        default="md",
        help="Output format (default: md)",
    )
    p.add_argument("--api", default="http://127.0.0.1:8000", help="CoreAPI base URL")
    p.add_argument(
        "--endpoint",
        default=DEFAULT_ENDPOINT,
        help=f"SSE path (default: {DEFAULT_ENDPOINT})",
    )
    p.add_argument(
        "--language",
        choices=("en", "am", "om"),
        default=None,
        help="Force response language (default: auto)",
    )
    p.add_argument("--model", default=None, help="OpenRouter model id (optional)")
    p.add_argument(
        "--reasoning",
        action="store_true",
        help="Enable extended reasoning (slower/costlier)",
    )
    p.add_argument(
        "--timeout",
        type=float,
        default=DEFAULT_TIMEOUT_SECONDS,
        help=f"HTTP timeout seconds per request (default: {int(DEFAULT_TIMEOUT_SECONDS)})",
    )
    p.add_argument(
        "--limit",
        type=int,
        default=None,
        help="Only the first N questions (per-question mode)",
    )
    p.add_argument(
        "--full-dump",
        action="store_true",
        help="Legacy: send the entire exam in one request (not recommended)",
    )
    p.add_argument(
        "--start",
        type=int,
        default=1,
        help="1-based question number to start from (per-question mode)",
    )
    args = p.parse_args()

    if not args.exam.exists():
        raise SystemExit(f"File not found: {args.exam}")
    if args.exam.suffix.lower() not in SUPPORTED_SUFFIXES:
        raise SystemExit("Input must be a .doc or .docx file")

    exam_text = extract_exam_text(args.exam)
    out = args.output
    if out is None:
        out = args.exam.with_name(f"{args.exam.stem}_answers.{args.format}")
    meta_path = out.with_suffix(".meta.json")

    print(f"Extracted {len(exam_text)} chars from {args.exam.name}", file=sys.stderr)

    if args.full_dump:
        print(
            f"FULL-DUMP mode: one call to {args.api}{args.endpoint} …",
            file=sys.stderr,
        )
        result = stream_exam(
            api=args.api,
            endpoint=args.endpoint,
            message=PROMPT + exam_text,
            language=args.language,
            model=args.model,
            enable_reasoning=args.reasoning,
            timeout=args.timeout,
        )
        answer = result["answer"]
        meta = {"mode": "full_dump", **result}
        if not answer:
            raise SystemExit("Empty answer from agent")
    else:
        jobs = build_question_messages(exam_text)
        if args.start > 1:
            jobs = [j for j in jobs if j["n"] >= args.start]
        if args.limit is not None:
            jobs = jobs[: args.limit]
        if not jobs:
            raise SystemExit("No questions parsed from exam after stripping keys")
        print(
            f"Per-question mode: {len(jobs)} calls to {args.api}{args.endpoint} …",
            file=sys.stderr,
        )
        parts: list[str] = []
        meta_items: list[dict] = []
        for job in jobs:
            print(f"\n--- Q{job['n']} ---", file=sys.stderr)
            result = stream_exam(
                api=args.api,
                endpoint=args.endpoint,
                message=job["message"],
                language=args.language,
                model=args.model,
                enable_reasoning=args.reasoning,
                timeout=args.timeout,
            )
            ans = result["answer"]
            if not ans:
                raise SystemExit(f"Empty answer from agent for Q{job['n']}")
            parts.append(f"## Q{job['n']}\n{ans}")
            meta_items.append(
                {
                    "n": job["n"],
                    "stem": job["stem"],
                    "thread_id": result.get("thread_id"),
                    "answer": ans,
                    "citations": result.get("citations"),
                    "grounding": result.get("grounding"),
                    "tool_queries": result.get("tool_queries") or [],
                }
            )
        answer = "\n\n".join(parts)
        meta = {"mode": "per_question", "items": meta_items}

    if args.format == "txt":
        body = answer.replace("## ", "").replace("# ", "")
    else:
        body = answer if answer.lstrip().startswith("#") else f"# Exam answers\n\n{answer}"

    out.write_text(body + "\n", encoding="utf-8")
    meta_path.write_text(json.dumps(meta, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(f"Wrote {out}", file=sys.stderr)
    print(f"Wrote {meta_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
