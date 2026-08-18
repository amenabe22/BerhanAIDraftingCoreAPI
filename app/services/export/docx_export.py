"""Markdown-ish text → DOCX bytes (python-docx)."""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.services.export.md_lines import classify_md_line, heading_text

INK = RGBColor(0x00, 0x00, 0x00)


def _strip_md(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def _set_run(run, *, size: float, bold: bool, color: RGBColor) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _add_bottom_border(paragraph, color_hex: str = "000000", size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def markdown_to_docx_bytes(markdown: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    for raw in (markdown or "").split("\n"):
        kind = classify_md_line(raw)
        if kind in {"blank", "rule"}:
            continue

        if kind == "title":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(10)
            run = para.add_run(_strip_md(heading_text(raw, kind)))
            _set_run(run, size=22, bold=True, color=INK)
            _add_bottom_border(para)
        elif kind == "h2":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(8)
            run = para.add_run(_strip_md(heading_text(raw, kind)))
            _set_run(run, size=16, bold=True, color=INK)
        elif kind == "h3":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(_strip_md(heading_text(raw, kind)))
            _set_run(run, size=13, bold=True, color=INK)
        elif kind == "bullet":
            para = doc.add_paragraph(_strip_md(heading_text(raw, kind)), style="List Bullet")
            para.paragraph_format.space_after = Pt(4)
            for run in para.runs:
                _set_run(run, size=10.5, bold=False, color=INK)
        elif kind == "ordered":
            para = doc.add_paragraph(
                _strip_md(heading_text(raw, kind)), style="List Number"
            )
            para.paragraph_format.space_after = Pt(4)
            for run in para.runs:
                _set_run(run, size=10.5, bold=False, color=INK)
        else:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.line_spacing = 1.15
            run = para.add_run(_strip_md(raw.strip()))
            _set_run(run, size=10.5, bold=False, color=INK)

    if len(doc.paragraphs) == 0:
        doc.add_paragraph("(Empty document)")

    buffer = io.BytesIO()
    doc.save(buffer)
    data = buffer.getvalue()
    buffer.close()
    return data
