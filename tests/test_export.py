"""Unit tests for TipTap → Contabo PDF/DOCX export."""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from fastapi.testclient import TestClient

from app.main import app
from app.services.export.contabo import (
    ContaboNotConfiguredError,
    contabo_configured,
    reset_client_cache,
    upload_bytes,
)
from app.services.export.docx_export import markdown_to_docx_bytes
from app.services.export.pdf import (
    INK,
    build_paragraph_styles,
    markdown_to_pdf_bytes,
)
from app.services.export.service import _safe_basename, export_document
from app.services.export.tiptap_text import tiptap_to_markdown, validate_tiptap_document

SAMPLE_DOC = {
    "type": "doc",
    "content": [
        {
            "type": "page",
            "content": [
                {
                    "type": "heading",
                    "attrs": {"level": 1},
                    "content": [{"type": "text", "text": "NDA"}],
                },
                {
                    "type": "paragraph",
                    "content": [
                        {
                            "type": "text",
                            "text": "This agreement is between Acme and Beta.",
                        }
                    ],
                },
                {
                    "type": "bulletList",
                    "content": [
                        {
                            "type": "listItem",
                            "content": [
                                {
                                    "type": "paragraph",
                                    "content": [{"type": "text", "text": "Confidentiality"}],
                                }
                            ],
                        }
                    ],
                },
            ],
        }
    ],
}

RICH_DOC = {
    "type": "doc",
    "content": [
        {
            "type": "page",
            "content": [
                {
                    "type": "heading",
                    "attrs": {"level": 2},
                    "content": [{"type": "text", "text": "Obligations"}],
                },
                {
                    "type": "orderedList",
                    "content": [
                        {
                            "type": "listItem",
                            "content": [
                                {
                                    "type": "paragraph",
                                    "content": [{"type": "text", "text": "Keep secrets"}],
                                }
                            ],
                        },
                        {
                            "type": "listItem",
                            "content": [
                                {
                                    "type": "paragraph",
                                    "content": [{"type": "text", "text": "Return materials"}],
                                }
                            ],
                        },
                    ],
                },
                {
                    "type": "blockquote",
                    "content": [
                        {
                            "type": "paragraph",
                            "content": [{"type": "text", "text": "Quoted clause"}],
                        }
                    ],
                },
            ],
        },
        {"type": "pageBreak"},
        {
            "type": "page",
            "content": [
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "Page two body"}],
                }
            ],
        },
    ],
}

AMHARIC_DOC = {
    "type": "doc",
    "content": [
        {
            "type": "page",
            "content": [
                {
                    "type": "heading",
                    "attrs": {"level": 1},
                    "content": [{"type": "text", "text": "ስምምነት"}],
                },
                {
                    "type": "paragraph",
                    "content": [{"type": "text", "text": "ይህ ስምምነት በሁለት ወገኖች መካከል ነው።"}],
                },
            ],
        }
    ],
}

EMPTY_TEXT_DOC = {
    "type": "doc",
    "content": [
        {
            "type": "page",
            "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "   "}]},
            ],
        }
    ],
}


@pytest.fixture
def client():
    return TestClient(app)


@pytest.fixture
def mock_upload():
    with (
        patch("app.services.export.service.contabo_configured", return_value=True),
        patch(
            "app.services.export.service.upload_bytes",
            side_effect=lambda key, data, content_type: f"https://cdn.test/{key}",
        ) as upload,
    ):
        yield upload


# ── OpenAPI / routing ─────────────────────────────────────────────────────────


def test_export_route_in_openapi(client: TestClient):
    paths = client.app.openapi()["paths"]
    assert "/drafting/export" in paths
    assert "/drafting/generate/stream" in paths


# ── TipTap validation ─────────────────────────────────────────────────────────


def test_validate_tiptap_ok():
    validate_tiptap_document(SAMPLE_DOC)  # no raise


@pytest.mark.parametrize(
    "doc",
    [
        {"type": "paragraph"},
        {"type": "doc", "content": []},
        {"type": "doc"},
        "not-a-dict",
        None,
        {"type": "doc", "content": "oops"},
    ],
)
def test_validate_tiptap_rejects_bad(doc):
    with pytest.raises(ValueError):
        validate_tiptap_document(doc)


# ── TipTap → markdown ─────────────────────────────────────────────────────────


def test_tiptap_to_markdown_basic():
    md = tiptap_to_markdown(SAMPLE_DOC)
    assert "# NDA" in md
    assert "Acme and Beta" in md
    assert "- Confidentiality" in md


def test_tiptap_to_markdown_rich():
    md = tiptap_to_markdown(RICH_DOC)
    assert "## Obligations" in md
    assert "1. Keep secrets" in md
    assert "2. Return materials" in md
    assert "> Quoted clause" in md
    assert "---" in md
    assert "Page two body" in md


def test_tiptap_to_markdown_amharic():
    md = tiptap_to_markdown(AMHARIC_DOC)
    assert "ስምምነት" in md
    assert "ወገኖች" in md


def test_tiptap_to_markdown_rejects_non_object():
    with pytest.raises(ValueError):
        tiptap_to_markdown([])  # type: ignore[arg-type]


# ── PDF / DOCX converters ─────────────────────────────────────────────────────


def test_markdown_to_pdf_bytes_is_pdf():
    data = markdown_to_pdf_bytes("# Title\n\nHello world body text.")
    assert data.startswith(b"%PDF")
    assert len(data) > 100


def test_markdown_to_pdf_empty_still_valid():
    data = markdown_to_pdf_bytes("")
    assert data.startswith(b"%PDF")


def test_markdown_to_docx_bytes_is_zip():
    data = markdown_to_docx_bytes("# Title\n\n- item one\n1. numbered\n## Section\nBody")
    assert data[:2] == b"PK"  # DOCX is a ZIP
    assert len(data) > 1000


def test_markdown_to_docx_empty_still_valid():
    data = markdown_to_docx_bytes("")
    assert data[:2] == b"PK"


def test_docx_headings_are_larger_than_body():
    from io import BytesIO

    from docx import Document

    data = markdown_to_docx_bytes("# MoyAts NDA\n\n## 1. Purpose\n\nBody clause text.")
    doc = Document(BytesIO(data))
    title, section, body = doc.paragraphs[0], doc.paragraphs[1], doc.paragraphs[2]
    assert title.runs[0].font.size.pt >= 18
    assert section.runs[0].font.size.pt >= 13
    body_pt = body.runs[0].font.size.pt if body.runs[0].font.size else 11
    assert title.runs[0].font.size.pt > body_pt


def test_pdf_heading_styles_are_larger_than_body():
    styles = build_paragraph_styles("Helvetica")
    assert styles["title"].fontSize >= 18
    assert styles["h2"].fontSize >= 13
    assert styles["h2"].fontSize > styles["body"].fontSize
    assert styles["h2"].spaceBefore >= 12
    assert styles["title"].textColor == INK
    assert styles["body"].textColor == INK


def test_numbered_clause_titles_classify_as_headings():
    from app.services.export.md_lines import classify_md_line

    assert classify_md_line("1. Purpose") == "h2"
    assert classify_md_line("## 1. Definitions") == "h2"
    assert classify_md_line("1. The receiving party shall keep secrets.") == "ordered"


def test_english_pdf_embeds_noto_sans():
    pdf = markdown_to_pdf_bytes("# MoyAts NDA\n\n## 1. Purpose\n\nBody clause text.")
    assert b"NotoSans" in pdf


def test_docx_numbered_section_title_is_large():
    from io import BytesIO

    from docx import Document

    data = markdown_to_docx_bytes("# MoyAts NDA\n\n1. Purpose\n\nBody clause text.")
    doc = Document(BytesIO(data))
    assert doc.paragraphs[1].runs[0].font.size.pt >= 13


def test_pdf_and_docx_handle_amharic():
    md = tiptap_to_markdown(AMHARIC_DOC)
    pdf = markdown_to_pdf_bytes(md)
    docx = markdown_to_docx_bytes(md)
    assert pdf.startswith(b"%PDF")
    assert docx[:2] == b"PK"


_MIXED_MD = (
    "# MoyAts NDA\n\n"
    "This agreement is between MoyAts PLC and Acme Trading PLC.\n\n"
    "ይህ ስምምነት በ MoyAts PLC መካከል ነው።\n"
)


def test_pdf_embeds_truetype_for_latin_and_ethiopic():
    """Helvetica-only PDFs render Ethiopic (and other Unicode) as black boxes."""
    pdf = markdown_to_pdf_bytes(_MIXED_MD)
    assert pdf.startswith(b"%PDF")
    assert b"/FontFile2" in pdf


def test_pdf_uses_ethiopic_lessan_for_amharic():
    pdf = markdown_to_pdf_bytes(_MIXED_MD)
    assert b"EthiopicLessan" in pdf


def test_pdf_text_extracts_amharic_and_latin():
    if not shutil.which("pdftotext"):
        pytest.skip("pdftotext is required to assert extractable glyphs")
    pdf = markdown_to_pdf_bytes(_MIXED_MD)
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as fh:
        fh.write(pdf)
        path = fh.name
    try:
        proc = subprocess.run(
            ["pdftotext", "-enc", "UTF-8", path, "-"],
            check=True,
            capture_output=True,
        )
    finally:
        Path(path).unlink(missing_ok=True)
    text = proc.stdout.decode("utf-8")
    assert "MoyAts NDA" in text
    assert "MoyAts PLC" in text
    assert "ስምምነት" in text
    assert "■" not in text


# ── Filename sanitization ─────────────────────────────────────────────────────


@pytest.mark.parametrize(
    ("raw", "expected"),
    [
        (None, "document"),
        ("", "document"),
        ("  ", "document"),
        ("nda-acme", "nda-acme"),
        ("My NDA!!!.pdf", "My_NDA_"),
        ("report.docx", "report"),
        ("a" * 200, "a" * 120),
    ],
)
def test_safe_basename(raw, expected):
    assert _safe_basename(raw) == expected


# ── Contabo client ────────────────────────────────────────────────────────────


def test_contabo_configured_false_when_bucket_missing():
    with patch("app.services.export.contabo.settings") as s:
        s.S3_ENDPOINT_URL = "https://eu2.contabostorage.com"
        s.S3_ACCESS_KEY_ID = "key"
        s.S3_SECRET_ACCESS_KEY = "secret"
        s.S3_BUCKET_NAME = ""
        assert contabo_configured() is False


def test_contabo_configured_true_when_all_set():
    with patch("app.services.export.contabo.settings") as s:
        s.S3_ENDPOINT_URL = "https://eu2.contabostorage.com"
        s.S3_ACCESS_KEY_ID = "key"
        s.S3_SECRET_ACCESS_KEY = "secret"
        s.S3_BUCKET_NAME = "berhan-core"
        assert contabo_configured() is True


def test_upload_bytes_presigned_url():
    reset_client_cache()
    mock_client = MagicMock()
    mock_client.generate_presigned_url.return_value = "https://presigned.example/obj"

    with (
        patch("app.services.export.contabo.contabo_configured", return_value=True),
        patch("app.services.export.contabo._s3_client", return_value=mock_client),
        patch("app.services.export.contabo.settings") as s,
    ):
        s.S3_BUCKET_NAME = "berhan-core"
        s.S3_PUBLIC_BASE_URL = None
        s.S3_PRESIGN_EXPIRY_SECONDS = 3600
        url = upload_bytes("exports/x.pdf", b"%PDF-1.4", "application/pdf")

    mock_client.put_object.assert_called_once()
    assert url == "https://presigned.example/obj"
    mock_client.generate_presigned_url.assert_called_once()


def test_upload_bytes_public_base_url():
    reset_client_cache()
    mock_client = MagicMock()

    with (
        patch("app.services.export.contabo.contabo_configured", return_value=True),
        patch("app.services.export.contabo._s3_client", return_value=mock_client),
        patch("app.services.export.contabo.settings") as s,
    ):
        s.S3_BUCKET_NAME = "berhan-core"
        s.S3_PUBLIC_BASE_URL = "https://cdn.example/berhan-core"
        s.S3_PRESIGN_EXPIRY_SECONDS = 3600
        url = upload_bytes("exports/x.pdf", b"%PDF-1.4", "application/pdf")

    assert url == "https://cdn.example/berhan-core/exports/x.pdf"
    mock_client.generate_presigned_url.assert_not_called()


def test_s3_client_raises_when_not_configured():
    reset_client_cache()
    with patch("app.services.export.contabo.contabo_configured", return_value=False):
        from app.services.export import contabo as contabo_mod

        with pytest.raises(ContaboNotConfiguredError):
            contabo_mod._s3_client()


# ── export_document service ───────────────────────────────────────────────────


def test_export_service_both_formats(mock_upload):
    result = export_document(SAMPLE_DOC, formats=["pdf", "docx"], filename="nda-test")
    assert result["filename"] == "nda-test"
    assert result["pdf_url"].startswith("https://cdn.test/")
    assert result["docx_url"].startswith("https://cdn.test/")
    assert result["keys"]["pdf"].endswith(".pdf")
    assert result["keys"]["docx"].endswith(".docx")
    assert "exports/" in result["keys"]["pdf"]
    assert mock_upload.call_count == 2


def test_export_service_pdf_only(mock_upload):
    result = export_document(SAMPLE_DOC, formats=["pdf"])
    assert result["pdf_url"]
    assert result["docx_url"] is None
    assert list(result["keys"]) == ["pdf"]
    assert mock_upload.call_count == 1


def test_export_service_docx_only(mock_upload):
    result = export_document(SAMPLE_DOC, formats=["docx"])
    assert result["docx_url"]
    assert result["pdf_url"] is None
    assert list(result["keys"]) == ["docx"]


def test_export_service_default_formats_both(mock_upload):
    result = export_document(SAMPLE_DOC)
    assert result["pdf_url"] and result["docx_url"]
    assert mock_upload.call_count == 2


def test_export_service_strips_unknown_formats(mock_upload):
    # unknown formats filtered; only pdf remains
    result = export_document(SAMPLE_DOC, formats=["pdf", "rtf"])  # type: ignore[list-item]
    assert result["pdf_url"]
    assert result["docx_url"] is None


def test_export_service_rejects_empty_formats():
    with (
        patch("app.services.export.service.contabo_configured", return_value=True),
        pytest.raises(ValueError, match="formats must include"),
    ):
        export_document(SAMPLE_DOC, formats=[])


def test_export_service_rejects_only_unknown_formats():
    with (
        patch("app.services.export.service.contabo_configured", return_value=True),
        pytest.raises(ValueError, match="formats must include"),
    ):
        export_document(SAMPLE_DOC, formats=["rtf"])  # type: ignore[list-item]


def test_export_service_rejects_empty_text():
    with (
        patch("app.services.export.service.contabo_configured", return_value=True),
        pytest.raises(ValueError, match="no exportable text"),
    ):
        export_document(EMPTY_TEXT_DOC, formats=["pdf"])


def test_export_service_requires_contabo():
    with (
        patch("app.services.export.service.contabo_configured", return_value=False),
        pytest.raises(ContaboNotConfiguredError),
    ):
        export_document(SAMPLE_DOC, formats=["pdf"])


def test_export_service_upload_payload_types(mock_upload):
    export_document(SAMPLE_DOC, formats=["pdf", "docx"])
    types = {c.kwargs.get("content_type") or c.args[2] for c in mock_upload.call_args_list}
    assert "application/pdf" in types
    assert (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in types
    )


# ── HTTP endpoint ─────────────────────────────────────────────────────────────


def test_export_endpoint_requires_document(client: TestClient):
    r = client.post("/drafting/export", json={})
    assert r.status_code == 422


def test_export_endpoint_503_when_unconfigured(client: TestClient):
    with patch("app.services.export.service.contabo_configured", return_value=False):
        r = client.post(
            "/drafting/export",
            json={"document": SAMPLE_DOC, "formats": ["pdf"]},
        )
    assert r.status_code == 503
    assert "Contabo" in r.json()["detail"] or "S3" in r.json()["detail"]


def test_export_endpoint_400_invalid_document(client: TestClient, mock_upload):
    r = client.post(
        "/drafting/export",
        json={"document": {"type": "doc", "content": []}, "formats": ["pdf"]},
    )
    assert r.status_code == 400


def test_export_endpoint_400_empty_text(client: TestClient, mock_upload):
    r = client.post(
        "/drafting/export",
        json={"document": EMPTY_TEXT_DOC, "formats": ["pdf"]},
    )
    assert r.status_code == 400


def test_export_endpoint_200_pdf_only(client: TestClient, mock_upload):
    r = client.post(
        "/drafting/export",
        json={"document": SAMPLE_DOC, "formats": ["pdf"], "filename": "sample"},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["filename"] == "sample"
    assert body["pdf_url"].startswith("https://cdn.test/")
    assert body["docx_url"] is None
    assert "pdf" in body["keys"]


def test_export_endpoint_200_both_formats(client: TestClient, mock_upload):
    r = client.post(
        "/drafting/export",
        json={"document": SAMPLE_DOC, "formats": ["pdf", "docx"], "filename": "full"},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["pdf_url"] and body["docx_url"]
    assert set(body["keys"]) == {"pdf", "docx"}


def test_export_endpoint_200_amharic(client: TestClient, mock_upload):
    r = client.post(
        "/drafting/export",
        json={"document": AMHARIC_DOC, "formats": ["pdf", "docx"]},
    )
    assert r.status_code == 200
    assert r.json()["pdf_url"] and r.json()["docx_url"]


def test_export_endpoint_500_on_upload_failure(client: TestClient):
    with (
        patch("app.services.export.service.contabo_configured", return_value=True),
        patch(
            "app.services.export.service.upload_bytes",
            side_effect=RuntimeError("S3 put failed"),
        ),
    ):
        r = client.post(
            "/drafting/export",
            json={"document": SAMPLE_DOC, "formats": ["pdf"]},
        )
    assert r.status_code == 500
    assert "Export failed" in r.json()["detail"]
