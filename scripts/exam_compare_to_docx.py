#!/usr/bin/env python3
"""Write a formatted Word report of Berhan answers vs exam-sheet keys."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

EXMS = ROOT / "exms"
COMPARE_JSON = Path("/tmp/exam_compare.json")
MISMATCH_JSON = EXMS / "mismatch_report.json"
OUT_PATH = EXMS / "exam_answer_key_compare.docx"

def rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


NAVY = "1B365D"
NAVY_RGB = rgb(NAVY)
WHITE = "FFFFFF"
GREEN = "C6EFCE"
RED = "FFC7CE"
ORANGE = "FCE4D6"
GRAY = "F2F2F2"
GRAY_DK = rgb("595959")
RED_DK = rgb("9C0006")
GOLD_BG = "FFF2CC"
ROW_ALT = "F7F9FC"

STATUS_FILL = {
    "correct": GREEN,
    "wrong": RED,
    "kb-miss": ORANGE,
    "no_key": GRAY,
    "unanswered": GRAY,
}
STATUS_LABEL = {
    "correct": "Correct",
    "wrong": "Wrong",
    "kb-miss": "KB miss",
    "no_key": "No key",
    "unanswered": "Blank",
}

QHEAD_RE = re.compile(r"(?m)^##\s*Q(\d+)\s*$")
EXAM_IDS = {
    "Assosa University- Law Model Exit Exam - 2015 E.C.": "assosa-2015",
    "DDU LAW MODEL EXAM for Practice": "ddu-practice",
    "Debre Berhan University Model Exit Exam (2023)": "debre-berhan-2023",
    "Debre Markos University Model Exam": "debre-markos-university-model-exam",
    "HU- COL-MOCK Exam-2026": "hu-col-mock-2026",
    "LL.B Model Exam Debre Markos University": "llb-debre-markos",
    "LL.B Model Exit-Exam-Ambo Uni sol": "ll-b-model-exit-exam-ambo-uni-sol",
    "Samara University Law model exam": "samara",
    "UoG, SOL Model  Exam": "uog-sol",
    "Wolaita Sodo MODEL Exam 2018": "wolaita-sodo-2018",
    "Model_Exam_for_WSU,_SOL_2015_E_C_": "wsu-2015",
    "Model": "model-generic",
}
KB_MISS_RE = re.compile(r"(?i)knowledge base does not contain")


def exam_id_for_file(name: str) -> str:
    stem = Path(name).stem.rstrip(".")
    matches: list[tuple[int, str]] = []
    for prefix, eid in EXAM_IDS.items():
        p = prefix.rstrip(".")
        if stem == p or stem.startswith(p):
            matches.append((len(p), eid))
    if matches:
        matches.sort(reverse=True)
        return matches[0][1]
    return re.sub(r"[^a-z0-9]+", "-", stem.lower()).strip("-")[:40]


def parse_berhan(md_path: Path) -> dict[int, dict]:
    text = md_path.read_text(encoding="utf-8", errors="replace")
    parts = re.split(r"(?m)^##\s*Q(\d+)\s*$", text)
    out: dict[int, dict] = {}
    for i in range(1, len(parts), 2):
        n = int(parts[i])
        if n in out:
            continue
        body = parts[i + 1].strip()
        first = body.split("\n", 1)[0].strip() if body else ""
        out[n] = {"first": first, "body": body}
    return out


def find_answers_md(src_name: str) -> Path | None:
    stem = Path(src_name).stem
    for cand in (
        EXMS / f"{stem}_answers.md",
        EXMS / f"{stem.rstrip('.')}_answers.md",
        EXMS / f"{stem}._answers.md",
        EXMS / f"{stem.rstrip('.')}._answers.md",
    ):
        if cand.exists():
            return cand
    key = stem[:28]
    hits = sorted(EXMS.glob(f"{key}*_answers.md"))
    return hits[0] if len(hits) == 1 else None


def set_run_font(run, *, size: int = 10, bold: bool = False, color: RGBColor | None = None, name: str = "Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = "BFBFBF") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    color: RGBColor | None = None,
    fill: str | None = None,
    align: str = "left",
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text if text is not None else "")
    set_run_font(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    cell.vertical_alignment = 1  # center


def set_col_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    total = sum(widths_cm)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(total * 567)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for child in list(grid):
            grid.remove(child)
        for w in widths_cm:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(int(w * 567)))
            grid.append(gc)
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def repeat_header(table) -> None:
    tr = table.rows[0]._tr
    tr_pr = tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    tr_pr.append(hdr)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY_RGB
        run.font.name = "Calibri"


def add_body(doc: Document, text: str, *, italic: bool = False, size: int = 11, space_after: int = 8) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=NAVY_RGB if not italic else GRAY_DK)
    run.italic = italic


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld1, instr, fld2])
    set_run_font(run, size=9, color=GRAY_DK)


def load_stems(src_name: str) -> dict[int, str]:
    stem = Path(src_name).stem
    src = None
    for ext in (".docx", ".doc"):
        for cand in (EXMS / f"{stem}{ext}", EXMS / f"{stem.rstrip('.')}{ext}"):
            if cand.exists():
                src = cand
                break
        if src:
            break
    if src is None or src.suffix.lower() != ".docx":
        return {}
    try:
        from examtest import extract_docx_text
        from app.services.legal.exam_parse import split_exam_questions, strip_printed_keys

        raw = extract_docx_text(src)
        questions = split_exam_questions(strip_printed_keys(raw))
        return {q["n"]: q.get("stem") or "" for q in questions}
    except Exception:
        return {}


def add_summary_table(doc: Document, exams: list[dict]) -> None:
    headers = ["Exam", "Key source", "Scored", "Correct", "Wrong", "KB miss", "No key", "Score"]
    widths = [5.4, 2.4, 1.4, 1.5, 1.4, 1.5, 1.4, 1.4]
    table = doc.add_table(rows=1 + len(exams), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(255, 255, 255), fill=NAVY, align="center")
    for r, exam in enumerate(exams, start=1):
        pct = exam.get("pct")
        score = f"{pct:.1f}%" if isinstance(pct, (int, float)) else "—"
        fill = ROW_ALT if r % 2 == 0 else WHITE
        if exam.get("key_visible_in_prompt"):
            score = f"{score} *"
        values = [
            exam["exam"],
            exam.get("method") or "—",
            str(exam.get("scored") or 0),
            str(exam.get("correct") or 0),
            str(exam.get("wrong") or 0),
            str(exam.get("kb_miss") or 0),
            str(exam.get("no_key") or 0),
            score,
        ]
        aligns = ["left", "center", "center", "center", "center", "center", "center", "center"]
        for i, (val, align) in enumerate(zip(values, aligns)):
            cell_fill = fill
            if i == 7 and isinstance(pct, (int, float)):
                cell_fill = GREEN if pct >= 80 else (ORANGE if pct >= 50 else RED)
            set_cell_text(table.rows[r].cells[i], val, size=8, fill=cell_fill, align=align, bold=(i in {0, 7}))
    set_col_widths(table, widths)
    repeat_header(table)


def add_cause_table(doc: Document, counts: dict) -> None:
    labels = {
        "option_mismatch": "Nearby option / EXCEPT miss",
        "prompt_conflict": "Guessed a letter, then KB-miss",
        "alignment_noise": "Question alignment off",
        "batch_retrieval": "Full-exam dump reused early retrieval",
        "kb_gap_real": "Topic not in statute KB",
        "wrong_instrument": "Wrong code family (e.g. Civil vs Commercial)",
        "key_in_prompt": "Printed key was in the prompt",
        "neighbor_article": "Neighboring article",
    }
    rows = sorted(counts.items(), key=lambda kv: -kv[1])
    table = doc.add_table(rows=1 + len(rows), cols=3)
    for i, h in enumerate(["Cause", "What it means", "Count"]):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(255, 255, 255), fill=NAVY, align="center")
    total = sum(counts.values()) or 1
    for r, (cause, n) in enumerate(rows, start=1):
        fill = ROW_ALT if r % 2 == 0 else WHITE
        set_cell_text(table.rows[r].cells[0], cause, size=9, fill=fill, bold=True)
        set_cell_text(table.rows[r].cells[1], labels.get(cause, cause), size=9, fill=fill)
        set_cell_text(table.rows[r].cells[2], f"{n}  ({100 * n / total:.0f}%)", size=9, fill=fill, align="center")
    set_col_widths(table, [4.2, 10.2, 2.0])
    repeat_header(table)


def add_compare_table(doc: Document, rows: list[dict]) -> None:
    headers = ["Q", "Gold key", "Berhan answer", "Result"]
    widths = [1.2, 2.6, 10.4, 2.2]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=RGBColor(255, 255, 255), fill=NAVY, align="center")
    for r, row in enumerate(rows, start=1):
        status = row["status"]
        fill = STATUS_FILL.get(status, WHITE)
        set_cell_text(table.rows[r].cells[0], str(row["q"]), size=9, fill=fill, align="center", bold=True)
        set_cell_text(table.rows[r].cells[1], row.get("gold") or "—", size=8, fill=GOLD_BG if row.get("gold") else fill, align="center", bold=True)
        set_cell_text(table.rows[r].cells[2], row.get("berhan") or "—", size=8, fill=fill)
        set_cell_text(
            table.rows[r].cells[3],
            STATUS_LABEL.get(status, status),
            size=8,
            fill=fill,
            align="center",
            bold=True,
        )
    set_col_widths(table, widths)
    repeat_header(table)


def add_mismatch_block(doc: Document, items: list[dict], berhan_map: dict[int, dict]) -> None:
    if not items:
        add_body(doc, "No classified mismatches for this exam.", italic=True, size=10)
        return
    for item in items:
        q = item["q"]
        heading = f"Q{q}  —  gold {item.get('gold') or '—'}  vs  Berhan  ·  {item.get('cause')}"
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(heading)
        set_run_font(run, size=11, bold=True, color=RGBColor(0x9C, 0x00, 0x06))

        if item.get("stem"):
            add_body(doc, f"Stem: {item['stem']}", italic=True, size=10, space_after=4)

        body = (berhan_map.get(q) or {}).get("body") or item.get("berhan") or ""
        if len(body) > 2500:
            body = body[:2500].rstrip() + " …"
        add_body(doc, f"Berhan: {body}", size=10, space_after=4)
        if item.get("why"):
            add_body(doc, f"Why: {item['why']}", size=10, space_after=6)


def build() -> Path:
    compare = json.loads(COMPARE_JSON.read_text(encoding="utf-8"))
    mismatch = {"cause_counts": {}, "mismatches": []}
    if MISMATCH_JSON.exists():
        mismatch = json.loads(MISMATCH_JSON.read_text(encoding="utf-8"))

    by_exam_mm: dict[str, list[dict]] = {}
    for row in mismatch.get("mismatches") or []:
        by_exam_mm.setdefault(row["exam_id"], []).append(row)

    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        header = section.header
        header.paragraphs[0].text = ""
        hp = header.paragraphs[0]
        run = hp.add_run("Berhan legal-search  ·  exam answer-key comparison")
        set_run_font(run, size=9, color=GRAY_DK)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fr = fp.add_run("Page ")
        set_run_font(fr, size=9, color=GRAY_DK)
        add_page_number(fp)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    trun = title.add_run("Exam answer-key comparison")
    set_run_font(trun, size=22, bold=True, color=NAVY_RGB)

    summary = compare.get("summary") or {}
    scored = summary.get("scored") or 0
    correct = summary.get("correct") or 0
    pct = summary.get("pct")
    add_body(
        doc,
        f"Overall: {correct} / {scored} scored items  ({pct}% correct) across "
        f"{summary.get('exams_scored')} of {summary.get('exams_total')} exams that had a recoverable key.",
        size=12,
    )
    add_body(
        doc,
        "Gold keys come from the original sheets (purple/bold/yellow marks, “Answer:” / “Option” lines, "
        "or Question-ID suffixes). Berhan answers come from the *_answers.md files produced by examtest.py. "
        "Exams marked with * printed the key in the dumped exam text, so a high score is not a clean "
        "legal-skill signal. Debre Markos .doc, Ambo, WSU, and the generic Model sheet had no usable key.",
        size=10,
        space_after=12,
    )

    add_heading(doc, "Score by exam", 1)
    add_summary_table(doc, compare.get("exams") or [])
    add_body(doc, "* Key was visible in the prompt (printed Answer:/Option/Question-ID line).", italic=True, size=9, space_after=14)

    add_heading(doc, "Why mismatches happened", 1)
    add_cause_table(doc, mismatch.get("cause_counts") or {})
    add_body(
        doc,
        "Cause counts cover wrong and knowledge-base-miss items only. Perfect printed-key exams "
        "(DDU, HU COL MOCK, UoG) therefore barely appear here.",
        italic=True,
        size=9,
        space_after=12,
    )

    for exam in compare.get("exams") or []:
        doc.add_page_break()
        add_heading(doc, exam["exam"], 1)
        pct = exam.get("pct")
        score_line = f"{exam.get('correct', 0)}/{exam.get('scored', 0)}"
        if isinstance(pct, (int, float)):
            score_line += f"  ({pct}%)"
        else:
            score_line += "  (not scored)"
        bits = [
            score_line,
            f"key method: {exam.get('method') or 'none'}",
            f"Berhan questions: {exam.get('berhan_n')}",
            f"gold keys: {exam.get('gold_n')}",
        ]
        if exam.get("key_visible_in_prompt"):
            bits.append("printed key was in the prompt")
        add_body(doc, "  ·  ".join(bits), size=11)
        if exam.get("note"):
            add_body(doc, exam["note"], italic=True, size=10)

        md = find_answers_md(exam["file"])
        berhan_map = parse_berhan(md) if md else {}
        details = {d["q"]: d for d in (exam.get("details") or [])}
        stems = load_stems(exam["file"])
        eid = exam_id_for_file(exam["file"])
        mm_rows = {r["q"]: r for r in by_exam_mm.get(eid, [])}

        qns = sorted(set(berhan_map) | set(details) | set(mm_rows))
        rows = []
        omitted_nokey = 0
        scored_exam = bool(exam.get("scored"))
        for qn in qns:
            det = details.get(qn) or {}
            rec = berhan_map.get(qn) or {}
            mm = mm_rows.get(qn) or {}
            body = rec.get("body") or ""
            status = det.get("status")
            if not status:
                if KB_MISS_RE.search(body):
                    status = "kb-miss"
                elif mm:
                    status = "wrong"
                else:
                    status = "no_key"
            elif status not in {"correct", "wrong"} and KB_MISS_RE.search(body):
                status = "kb-miss"
            gold = det.get("gold") or mm.get("gold") or ""
            berhan = rec.get("first") or det.get("berhan") or mm.get("berhan") or ""
            if scored_exam and status == "no_key" and not gold:
                omitted_nokey += 1
                continue
            rows.append({"q": qn, "gold": gold, "berhan": berhan, "status": status})

        add_heading(doc, "Question-by-question", 2)
        if rows:
            add_compare_table(doc, rows)
            if omitted_nokey:
                add_body(
                    doc,
                    f"{omitted_nokey} further Berhan items had no recoverable gold key and are omitted from this table.",
                    italic=True,
                    size=9,
                    space_after=8,
                )
        else:
            add_body(doc, "No paired questions for this exam.", italic=True)

        classified = sorted(by_exam_mm.get(eid, []), key=lambda r: r["q"])
        if classified:
            add_heading(doc, "Mismatches in detail", 2)
            for item in classified:
                if not item.get("stem") and item["q"] in stems:
                    item = {**item, "stem": stems[item["q"]]}
                add_mismatch_block(doc, [item], berhan_map)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size:,} bytes)")
